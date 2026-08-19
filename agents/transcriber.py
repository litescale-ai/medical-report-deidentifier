import os
import mimetypes
import pathlib
from pydantic import BaseModel, Field
from google.antigravity import Agent
from google.antigravity.types import Document, Image, Audio, Video
from utils.agent_config import build_agent_config

class VerbatimItem(BaseModel):
    timestamp_start: str = Field(description="The timestamp of the dialogue or activity start (e.g. '00:00:15' or 'page 1')")
    timestamp_end: str = Field(description="The timestamp of the dialogue or activity end (e.g. '00:00:25' or 'page 1')")
    speaker: str = Field(description="Speaker name or label (e.g. 'Interviewer', 'Subject', 'Narrator', 'Visual Action')")
    content: str = Field(description="Verbatim spoken text or descriptive activity text (if video action)")

class ExtractedTranscript(BaseModel):
    filename: str = Field(description="The source filename")
    file_type: str = Field(description="The type of file (e.g. 'PDF Document', 'Audio Interview', 'Video Assessment')")
    summary: str = Field(description="A brief high-level overview of this session/file")
    items: list[VerbatimItem] = Field(description="The detailed chronological verbatim transcription/extraction items")

def _read_text_file(filepath: str) -> str:
    """Read a plain-text file and return its contents as a string."""
    with open(filepath, "r", encoding="utf-8") as f:
        return f.read()

def _read_docx_file(filepath: str) -> str:
    """Extract text from a .docx file using python-docx.

    Falls back to raw XML tag-stripping if the DOCX has malformed internal
    XML that python-docx (lxml) cannot parse.
    """
    try:
        import docx
    except ImportError:
        raise ImportError("python-docx is required to process .docx files. Install it with: pip install python-docx")

    try:
        doc = docx.Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except Exception:
        # Fallback: extract text from the raw XML inside the ZIP
        return _read_docx_raw(filepath)


def _read_docx_raw(filepath: str) -> str:
    """Fallback text extraction for corrupted DOCX files.

    Opens the DOCX as a ZIP archive and strips XML tags from
    word/document.xml using regex, producing plain text.
    """
    import re
    import zipfile

    with zipfile.ZipFile(filepath) as z:
        parts = []
        for xml_name in ("word/document.xml", "word/header1.xml", "word/footer1.xml"):
            if xml_name in z.namelist():
                raw = z.read(xml_name).decode("utf-8", errors="replace")
                clean = re.sub(r"<[^>]+>", " ", raw)
                clean = re.sub(r"\s+", " ", clean).strip()
                if clean:
                    parts.append(clean)
    return "\n\n".join(parts)

# Extensions that should be read as plain text strings (not passed to Document)
_TEXT_EXTENSIONS = {".txt", ".md", ".markdown", ".rst", ".csv", ".tsv", ".json", ".log"}
_DOCX_EXTENSIONS = {".docx", ".doc"}

def load_media_file(filepath: str):
    """Loads a media file into the correct Antigravity media type."""
    ext = os.path.splitext(filepath)[1].lower()
    
    # Handle text-readable files directly as string
    if ext in _TEXT_EXTENSIONS:
        return _read_text_file(filepath)
    
    # Handle Word documents by extracting text
    if ext in _DOCX_EXTENSIONS:
        return _read_docx_file(filepath)
            
    mime_type, _ = mimetypes.guess_type(filepath)
    if not mime_type:
        # Fallback based on extension
        if ext in ['.pdf', '.rtf', '.html', '.xml']:
            return Document.from_file(filepath)
        elif ext in ['.png', '.jpg', '.jpeg', '.webp']:
            return Image.from_file(filepath)
        elif ext in ['.mp3', '.wav', '.m4a', '.flac']:
            return Audio.from_file(filepath)
        elif ext in ['.mp4', '.webm', '.mov', '.avi']:
            return Video.from_file(filepath)
        raise ValueError(f"Unknown extension and MIME type for file {filepath}")
        
    if mime_type.startswith("image/"):
        return Image.from_file(filepath)
    elif mime_type.startswith("audio/"):
        return Audio.from_file(filepath)
    elif mime_type.startswith("video/"):
        return Video.from_file(filepath)
    elif mime_type == "application/pdf":
        return Document.from_file(filepath)
    elif mime_type.startswith("text/"):
        # Read as plain text to avoid unsupported Document MIME types
        return _read_text_file(filepath)
    else:
        # Last resort: try reading as text
        try:
            return _read_text_file(filepath)
        except (UnicodeDecodeError, ValueError):
            return Document.from_file(filepath)

async def transcribe_media(filepath: str, api_key: str = None, backend: str = None, gemini_model: str = None, ollama_model: str = None) -> dict:
    """Uses the TranscriberAgent to extract verbatim text and descriptions from a media file."""
    filename = os.path.basename(filepath)
    media = load_media_file(filepath)
    
    system_instructions = (
        "You are an expert multimodal medical transcriber. Your task is to process input media files "
        "(audio, video, documents) and extract all spoken text, written text, and detailed visual descriptions "
        "(especially of subjects' physical/behavioral activities in video) verbatim.\n"
        "Ensure that for audio/video you capture timing indicators (e.g. '00:01:23'). "
        "For PDF/text documents, use page numbers or logical headers as timing/location indicators.\n"
        "Do not summarize or paraphrase dialogue - extract it word-for-word. "
        "For visual descriptions in video, log them chronologically as 'Visual Action' speaker items, "
        "describing the patient's behaviors, movements, or physical state in detail."
    )
    
    config = build_agent_config(
        system_instructions=system_instructions,
        response_schema=ExtractedTranscript,
        backend=backend,
        api_key=api_key,
        gemini_model=gemini_model,
        ollama_model=ollama_model,
    )
    
    prompt = (
        f"Verbatim extract and transcribe all content from this file '{filename}'.\n"
        "Provide verbatim dialogue and detailed behavioral activity descriptions where applicable."
    )
    
    async with Agent(config=config) as agent:
        # Pass both the text prompt and the loaded media content in the chat turn
        if isinstance(media, str):
            # If it was a plain text file, we just append the text
            full_prompt = f"{prompt}\n\n=== FILE CONTENT ===\n{media}"
            response = await agent.chat(full_prompt)
        else:
            response = await agent.chat([prompt, media])
            
        data = await response.structured_output()
        return data

