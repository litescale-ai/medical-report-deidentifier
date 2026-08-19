"""Document editor for in-place PDF and DOCX de-identification.

Provides format-preserving find-and-replace for:
  - PDF files (via PyMuPDF): redact + insert with font matching
  - DOCX files (via python-docx): run-level replacement preserving styles

Also provides a synthesis summary writer for the companion output file.
"""

import os
import json
from typing import Optional


# ---------------------------------------------------------------------------
# PDF Editing (PyMuPDF)
# ---------------------------------------------------------------------------

def _int_to_rgb(color_int: int) -> tuple:
    """Convert an integer colour (0xRRGGBB) to a (r, g, b) float tuple."""
    if isinstance(color_int, (list, tuple)):
        return tuple(color_int)
    r = ((color_int >> 16) & 0xFF) / 255.0
    g = ((color_int >> 8) & 0xFF) / 255.0
    b = (color_int & 0xFF) / 255.0
    return (r, g, b)


def _pymupdf_fontname_to_base(fontname: str) -> str:
    """Map a PDF-internal font name to a PyMuPDF base-14 font name.

    PyMuPDF's insert_text only accepts base-14 font names.  We detect
    bold and italic variants to preserve weight/style.

    Base-14 families:
      Helvetica:  helv, hebo (bold), heit (italic), hebi (bold-italic)
      Times:      tiro, tibo (bold), tiit (italic), tibi (bold-italic)
      Courier:    cour, cobo (bold), coit (italic), cobi (bold-italic)
    """
    fn = fontname.lower()

    # Detect weight/style flags
    is_bold = "bold" in fn or "black" in fn or "heavy" in fn
    is_italic = "italic" in fn or "oblique" in fn

    # Detect family
    if "courier" in fn or "mono" in fn:
        if is_bold and is_italic:
            return "cobi"
        if is_bold:
            return "cobo"
        if is_italic:
            return "coit"
        return "cour"

    if "times" in fn or "serif" in fn:
        if is_bold and is_italic:
            return "tibi"
        if is_bold:
            return "tibo"
        if is_italic:
            return "tiit"
        return "tiro"

    if "symbol" in fn:
        return "symb"
    if "zapf" in fn:
        return "zadb"

    # Default: Helvetica family
    if is_bold and is_italic:
        return "hebi"
    if is_bold:
        return "hebo"
    if is_italic:
        return "heit"
    return "helv"


def _extract_span_style(page, search_rect) -> dict:
    """Find the best text style (font, size, color) that intersects with the given search rectangle."""
    import pymupdf
    blocks = page.get_text("dict")["blocks"]
    
    best_span = None
    max_area = 0
    
    for b in blocks:
        if b.get("type", 0) != 0:
            continue
        for l in b.get("lines", []):
            for s in l.get("spans", []):
                span_rect = pymupdf.Rect(s["bbox"])
                intersect = span_rect.intersect(search_rect)
                area = intersect.get_area()
                if area > max_area:
                    max_area = area
                    best_span = s
                    
    if best_span:
        return {
            "fontname": best_span.get("font", "helv"),
            "size": best_span.get("size", 11.0),
            "color": _int_to_rgb(best_span.get("color", 0)),
            "flags": best_span.get("flags", 0),
        }
        
    # Fallback default
    return {
        "fontname": "helv",
        "size": 11.0,
        "color": (0, 0, 0),
        "flags": 0,
    }


def _is_pdf_searchable(doc, replacement_map: dict[str, str]) -> bool:
    """Quick check: does the PDF have a usable text layer?

    Returns True if search_for() can find at least one of the replacement
    targets anywhere in the document.  Returns False for scanned / CIDFont
    PDFs with no ToUnicode mapping.
    """
    sorted_keys = sorted(replacement_map.keys(), key=len, reverse=True)
    for page in doc:
        for target in sorted_keys:
            if not target.strip():
                continue
            if page.search_for(target):
                return True
    return False


def _ocr_pdf(input_path: str) -> str:
    """Run OCR on a PDF to produce a searchable copy.

    Uses ocrmypdf (Tesseract) with force_ocr to replace any broken
    text layer with a fresh OCR'd one.  Returns the path to the OCR'd
    temporary file.
    """
    import tempfile
    import ocrmypdf

    fd, ocr_path = tempfile.mkstemp(suffix=".pdf")
    os.close(fd)

    ocrmypdf.ocr(
        input_path,
        ocr_path,
        force_ocr=True,
        optimize=1,
        progress_bar=False,
    )
    return ocr_path


def _redact_pdf(input_path: str, output_path: str, replacement_map: dict[str, str]) -> int:
    """Core redaction logic: search → redact → insert replacement text.

    Returns the total number of replacements made.
    """
    import pymupdf

    doc = pymupdf.open(input_path)
    total_replacements = 0
    sorted_keys = sorted(replacement_map.keys(), key=len, reverse=True)

    for page in doc:
        replacements = []
        for target in sorted_keys:
            if not target.strip():
                continue
            rects = page.search_for(target)
            for rect in rects:
                style = _extract_span_style(page, rect)
                replacements.append((rect, replacement_map[target], style))

        if not replacements:
            continue

        total_replacements += len(replacements)

        # Redact all matched regions
        for rect, _, _ in replacements:
            page.add_redact_annot(rect, fill=(1, 1, 1))
        page.apply_redactions()

        # Insert replacement text at original positions
        for rect, replacement, style in replacements:
            base_font = _pymupdf_fontname_to_base(style["fontname"])
            fontsize = style["size"]
            color = style["color"]

            insert_point = pymupdf.Point(rect.x0, rect.y1 - 1)
            page.insert_text(
                insert_point,
                replacement,
                fontsize=fontsize,
                fontname=base_font,
                color=color,
            )

    doc.save(output_path)
    doc.close()
    return total_replacements


def deidentify_pdf(
    input_path: str,
    output_path: str,
    replacement_map: dict[str, str],
) -> bool:
    """Replace PII strings in a PDF with pseudonym hashes, preserving layout.

    Uses a two-phase approach:
      1. Try direct text search on the original PDF.
      2. If the text layer is unsearchable (scanned / broken CIDFonts),
         automatically OCR the PDF first, then retry redaction.

    Args:
        input_path: Path to the original PDF.
        output_path: Path to save the de-identified PDF.
        replacement_map: Mapping of real PII strings → pseudonym hashes.

    Returns:
        True if the PDF was successfully processed with at least one replacement.
    """
    import pymupdf

    # First attempt: direct redaction
    doc = pymupdf.open(input_path)
    searchable = _is_pdf_searchable(doc, replacement_map)
    doc.close()

    if searchable:
        count = _redact_pdf(input_path, output_path, replacement_map)
        return count > 0

    # Fallback: OCR the PDF, then redact the OCR'd version
    import logging
    logger = logging.getLogger(__name__)
    logger.info("PDF text layer is unsearchable — running OCR fallback...")

    ocr_path = None
    try:
        ocr_path = _ocr_pdf(input_path)
        count = _redact_pdf(ocr_path, output_path, replacement_map)
        if count > 0:
            logger.info(f"OCR fallback succeeded: {count} replacements made")
            return True
        else:
            logger.warning("OCR fallback produced 0 replacements — OCR may not have recognised the text")
            # Still copy the OCR'd version so the user at least gets a searchable PDF
            import shutil
            shutil.copy2(ocr_path, output_path)
            return False
    except Exception as e:
        logger.error(f"OCR fallback failed: {e}")
        # Copy original as-is so downstream doesn't break
        import shutil
        shutil.copy2(input_path, output_path)
        return False
    finally:
        if ocr_path and os.path.exists(ocr_path):
            os.unlink(ocr_path)


# ---------------------------------------------------------------------------
# DOCX Editing (python-docx)
# ---------------------------------------------------------------------------

def _replace_in_runs(paragraph, replacement_map: dict[str, str]) -> int:
    """Replace PII text in a paragraph's runs, preserving per-run formatting.

    Handles the common case where a PII string is contained within a single
    run.  For cross-run splits, falls back to a joined-run replacement
    strategy that preserves the formatting of the first matching run.
    """
    sorted_keys = sorted(replacement_map.keys(), key=len, reverse=True)
    replacements_made = 0

    # First pass: simple per-run replacement
    for run in paragraph.runs:
        for real_text in sorted_keys:
            if real_text in run.text:
                run.text = run.text.replace(real_text, replacement_map[real_text])
                replacements_made += 1

    # Second pass: handle cross-run splits
    full_text = paragraph.text
    for real_text in sorted_keys:
        if real_text not in full_text:
            continue

        # Check if it was already handled in per-run pass
        remaining = "".join(r.text for r in paragraph.runs)
        if real_text not in remaining:
            continue

        # Cross-run replacement: find the runs that span this text
        if _replace_across_runs(paragraph, real_text, replacement_map[real_text]):
            replacements_made += 1
            
    return replacements_made


def _replace_across_runs(paragraph, target: str, replacement: str) -> bool:
    """Replace text that spans multiple runs in a paragraph.

    Keeps the formatting of the first run that contains part of the target.
    Returns True if replaced.
    """
    runs = paragraph.runs
    if not runs:
        return False

    # Build a character-to-run mapping
    char_positions = []  # list of (run_index, char_index_in_run)
    for ri, run in enumerate(runs):
        for ci in range(len(run.text)):
            char_positions.append((ri, ci))

    full_text = "".join(r.text for r in runs)
    start_idx = full_text.find(target)
    if start_idx == -1:
        return False

    end_idx = start_idx + len(target)

    # Determine which runs are affected
    start_run, start_char = char_positions[start_idx]
    end_run, end_char = char_positions[end_idx - 1]

    # Put the replacement text into the first affected run
    runs[start_run].text = (
        runs[start_run].text[:start_char]
        + replacement
        + runs[end_run].text[end_char + 1:]
    )

    # Clear intermediate and end runs (if different from start)
    for ri in range(start_run + 1, end_run + 1):
        runs[ri].text = ""
        
    return True


def _process_paragraphs(paragraphs, replacement_map: dict[str, str]) -> int:
    """Apply replacements to a list of paragraphs. Returns number of replacements made."""
    count = 0
    for paragraph in paragraphs:
        count += _replace_in_runs(paragraph, replacement_map)
    return count


def deidentify_docx(
    input_path: str,
    output_path: str,
    replacement_map: dict[str, str],
) -> bool:
    """Replace PII strings in a DOCX with pseudonym hashes, preserving formatting.

    Processes body paragraphs, tables, headers, and footers.

    Args:
        input_path: Path to the original DOCX.
        output_path: Path to save the de-identified DOCX.
        replacement_map: Mapping of real PII strings → pseudonym hashes.

    Returns:
        True if the DOCX was successfully processed.
    """
    from docx import Document

    doc = Document(input_path)

    # Body paragraphs
    _process_paragraphs(doc.paragraphs, replacement_map)

    # Tables (including nested tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _process_paragraphs(cell.paragraphs, replacement_map)

    # Headers and footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header and header.is_linked_to_previous is False:
                _process_paragraphs(header.paragraphs, replacement_map)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            _process_paragraphs(cell.paragraphs, replacement_map)

        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer and footer.is_linked_to_previous is False:
                _process_paragraphs(footer.paragraphs, replacement_map)
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            _process_paragraphs(cell.paragraphs, replacement_map)

    doc.save(output_path)
    return True


# ---------------------------------------------------------------------------
# Synthesis Summary Writer
# ---------------------------------------------------------------------------

RECIPIENT_INSTRUCTIONS = """================================================================================
CRITICAL RECIPIENT INSTRUCTIONS - PLEASE READ CAREFULLY
================================================================================
This medical document has been pseudonymised for data privacy and security.
All Personal Identifiable Information (PII) including names of patients, clinicians,
relatives, facilities, and locations have been replaced with secure pseudonym hashes:
e.g., PATIENT_A4B3D2, DOCTOR_E8F9A0, etc.

IMPORTANT: You MUST preserve all these pseudonym hashes (e.g. PATIENT_XXXX) exactly 
as they appear in this document in any returned, updated, or generated reports.
Do NOT remove, edit, or replace these hashes. 

The originator retains the secure Identity Catalogue. When you return the processed 
report, the originator will use the preserved hashes to automatically and securely 
re-identify the patient and parties.
================================================================================

"""


def write_synthesis_summary(
    output_path: str,
    deidentified_chrono: dict,
    identity_catalogue: Optional[dict] = None,
) -> str:
    """Write a standalone synthesis summary file with recipient instructions.

    This companion file contains:
    - Recipient instructions banner
    - Patient synthesis summary
    - Identified categories
    - Pseudonym hash legend (entity types, no real names)

    Args:
        output_path: Path to save the synthesis summary.
        deidentified_chrono: The deidentified chronology data dict.
        identity_catalogue: Optional catalogue (only pseudonym types are included,
                           not real names — this file is shareable).

    Returns:
        The full text content of the synthesis summary.
    """
    summary = RECIPIENT_INSTRUCTIONS

    summary += f"PATIENT SYNTHESIS SUMMARY:\n{deidentified_chrono.get('patient_summary', '')}\n\n"

    categories = deidentified_chrono.get("categories_found", [])
    if categories:
        summary += "IDENTIFIED CATEGORIES:\n"
        for cat in categories:
            summary += f"- {cat}\n"
        summary += "\n"

    # Pseudonym legend (safe to share — no real names)
    if identity_catalogue:
        summary += "================================================================================\n"
        summary += "PSEUDONYM HASH LEGEND\n"
        summary += "================================================================================\n\n"
        for pseudonym_hash, details in identity_catalogue.items():
            entity_type = details.get("entity_type", "UNKNOWN")
            relationship = details.get("relationship_context", "")
            summary += f"  {pseudonym_hash}  [{entity_type}]  {relationship}\n"
        summary += "\n"

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(summary)

    return summary


# ---------------------------------------------------------------------------
# Dispatcher — choose editor by file extension
# ---------------------------------------------------------------------------

def deidentify_document(
    input_path: str,
    output_path: str,
    replacement_map: dict[str, str],
) -> bool:
    """Dispatch to the correct format-specific editor based on file extension.

    Args:
        input_path: Path to the original document.
        output_path: Path to save the de-identified document.
        replacement_map: Mapping of real PII strings → pseudonym hashes.

    Returns:
        True if the document was processed, False if format is unsupported.
    """
    ext = os.path.splitext(input_path)[1].lower()

    if ext == ".pdf":
        return deidentify_pdf(input_path, output_path, replacement_map)
    elif ext == ".docx":
        return deidentify_docx(input_path, output_path, replacement_map)
    else:
        return False


def reidentify_document(
    input_path: str,
    output_path: str,
    identity_catalogue: dict,
) -> bool:
    """Reverse de-identification on a PDF or DOCX document.

    Builds a reverse replacement map (pseudonym → canonical name) from the
    identity catalogue and applies it to the document.

    Args:
        input_path: Path to the pseudonymised document.
        output_path: Path to save the re-identified document.
        identity_catalogue: The secure identity catalogue mapping.

    Returns:
        True if the document was processed, False if format is unsupported.
    """
    ext = os.path.splitext(input_path)[1].lower()
    if ext not in (".pdf", ".docx"):
        return False

    # Build reverse map: pseudonym_hash → canonical_name
    reverse_map = {}
    for pseudonym_hash, details in identity_catalogue.items():
        reverse_map[pseudonym_hash] = details["canonical_name"]

    return deidentify_document(input_path, output_path, reverse_map)
